from fastapi import APIRouter, Depends, status, UploadFile, File, HTTPException
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.middleware.cors import CORSMiddleware
from fastapi import FastAPI
import os
from typing import List
import pdfplumber
from docx import Document
import shutil
from sqlalchemy.orm import Session  # Add this import

from backend.schemas import UserCreate, Token
from backend.database import get_db
from backend.controllers.user_controller import (
    signup_controller,
    login_controller,
    get_current_user_controller,
)
from backend.controllers.document_controller import upload_document_controller
from backend.models import User, Chat, Message  # Add Chat model import
from backend import models  # <-- Add this import at the top
from backend.utils.summarizer import get_text_summary

router = APIRouter()
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8501"],  # Streamlit default port
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@router.post("/signup", response_model=Token, status_code=status.HTTP_201_CREATED)
async def signup(user: UserCreate, db: Session = Depends(get_db)):
    return await signup_controller(user, db)

@router.post("/token", response_model=Token)
async def login(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    print(f"[DEBUG] Login attempt for user: {form_data.username}")
    try:
        result = await login_controller(form_data, db)
        print(f"[DEBUG] Login successful for user: {form_data.username}")
        return result
    except Exception as e:
        print(f"[DEBUG] Login failed for user {form_data.username}: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )

@router.get("/current_user")
async def read_current_user(current_user: User = Depends(get_current_user_controller)):
    # Ensure current_user is a valid User instance
    return {"username": current_user.username, "email": current_user.email}

@router.post("/upload_document")
async def upload_document(
    files: List[UploadFile] = File(...),
    current_user: User = Depends(get_current_user_controller)
):
    return await upload_document_controller(files, current_user)

@router.get("/user_chats")
async def get_user_chats(current_user: User = Depends(get_current_user_controller), db: Session = Depends(get_db)):
    """
    Return all chat threads (chat_ids) for the current user.
    """
    chats = db.query(models.Chat).filter(models.Chat.user_id == current_user.id).all()
    return {"chats": [{"chat_id": c.id, "name": c.name} for c in chats]}

@router.post("/create_chat", status_code=201)
async def create_chat(current_user: User = Depends(get_current_user_controller), db: Session = Depends(get_db)):
    """Create a new chat thread for the user."""
    try:
        print(f"[DEBUG] Creating new chat for user: {current_user.id}")
        chat = Chat(user_id=current_user.id, name="New Chat")  # Now Chat is imported
        db.add(chat)
        db.commit()  # Ensure the chat is committed
        db.refresh(chat)
        print(f"[DEBUG] Created chat with ID: {chat.id}")
        return {"chat_id": chat.id}
    except Exception as e:
        print(f"[ERROR] Failed to create chat: {e}")
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e)
        )

@router.get("/chat/{chat_id}/messages")
async def get_chat_messages(
    chat_id: int,
    page: int = 1,
    per_page: int = 20,
    current_user: User = Depends(get_current_user_controller),
    db: Session = Depends(get_db)
):
    """
    Return paginated messages for a chat, only if user owns the chat.
    Always return a valid response, even if there are no messages yet.
    """
    chat = db.query(models.Chat).filter(models.Chat.id == chat_id, models.Chat.user_id == current_user.id).first()
    if not chat:
        raise HTTPException(status_code=403, detail="Not authorized for this chat.")
    total_messages = db.query(models.Message).filter(models.Message.chat_id == chat_id).count()
    total_pages = max(1, (total_messages + per_page - 1) // per_page)
    messages = (
        db.query(models.Message)
        .filter(models.Message.chat_id == chat_id)
        .order_by(models.Message.id.asc())
        .offset((page - 1) * per_page)
        .limit(per_page)
        .all()
    )
    # Always return a valid response, even if messages is empty
    return {
        "messages": [{"role": m.role, "content": m.content} for m in messages],
        "total_pages": total_pages
    }

@router.post("/chat/{chat_id}/send_message", status_code=201)
async def send_message_to_chat(
    chat_id: int,
    data: dict,
    current_user: User = Depends(get_current_user_controller),
    db: Session = Depends(get_db)
):
    """
    Add a message to a chat, only if user owns the chat.
    Accepts both user and assistant messages.
    """
    chat = db.query(models.Chat).filter(models.Chat.id == chat_id, models.Chat.user_id == current_user.id).first()
    if not chat:
        raise HTTPException(status_code=403, detail="Not authorized for this chat.")
    from backend.models import Message
    role = data.get("role", "user")
    msg = Message(chat_id=chat_id, role=role, content=data.get("content", ""))
    db.add(msg)
    db.commit()
    db.refresh(msg)

    # Update chat name if this is the first user message and chat name is still "New Chat"
    if role == "user" and chat.name == "New Chat":
        user_msg_count = db.query(models.Message).filter(
            models.Message.chat_id == chat_id,
            models.Message.role == "user"
        ).count()
        if user_msg_count == 1:
            first_chars = data.get("content", "")[:10]
            if len(data.get("content", "")) > 10:
                first_chars += "..."
            chat.name = first_chars
            db.add(chat)
            db.commit()
            db.refresh(chat)

    return {"message_id": msg.id}

@router.delete("/chat/{chat_id}/delete", status_code=204)
async def delete_chat(
    chat_id: int,
    current_user: User = Depends(get_current_user_controller),
    db: Session = Depends(get_db)
):
    """
    Delete a chat and all its messages, only if user owns the chat.
    """
    chat = db.query(models.Chat).filter(models.Chat.id == chat_id, models.Chat.user_id == current_user.id).first()
    if not chat:
        raise HTTPException(status_code=403, detail="Not authorized for this chat.")
    db.query(models.Message).filter(models.Message.chat_id == chat_id).delete()
    db.delete(chat)
    db.commit()
    return

def convert_file_to_text(file_path: str, content_type: str) -> str:
    """Convert uploaded file to text based on file type."""
    try:
        # Handle PDF files
        if content_type == "application/pdf":
            all_text = []
            with pdfplumber.open(file_path) as pdf:
                # Extract text from all pages
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        all_text.append(text)

                # Extract tables from all pages
                for i, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    if tables:
                        all_text.append(f"\n[TABLE FROM PAGE {i+1}]")
                        for table in tables:
                            all_text.extend("\t".join(str(cell) for cell in row if cell) for row in table)
                
                return "\n".join(all_text)
        
        # Handle DOCX files
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(file_path)
            full_text = []
            
            # Extract headers and footers
            for section in doc.sections:
                header = section.header
                footer = section.footer
                if header:
                    full_text.append("[HEADER]")
                    for para in header.paragraphs:
                        full_text.append(para.text)
                if footer:
                    full_text.append("[FOOTER]")
                    for para in footer.paragraphs:
                        full_text.append(para.text)

            # Extract paragraphs
            full_text.append("[BODY]")
            for para in doc.paragraphs:
                full_text.append(para.text)

            # Extract tables
            for table in doc.tables:
                full_text.append("[TABLE]")
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    full_text.append("\t".join(row_data))

            return "\n".join(full_text)

        # Handle TXT files
        elif content_type == "text/plain":
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        else:
            return f"Unsupported file type: {content_type}"

    except Exception as e:
        return f"Error converting file: {str(e)}"

@router.post("/chat/{chat_id}/upload")
async def upload_chat_file(
    chat_id: int,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user_controller),
    db: Session = Depends(get_db)
):
    """Upload a file in the context of a chat conversation."""
    # Verify chat ownership
    chat = db.query(models.Chat).filter(models.Chat.id == chat_id, models.Chat.user_id == current_user.id).first()
    if not chat:
        raise HTTPException(status_code=403, detail="Not authorized for this chat.")

    # Determine the actual content type from filename
    file_ext = os.path.splitext(file.filename)[1].lower()
    content_type_map = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.txt': 'text/plain'
    }
    content_type = content_type_map.get(file_ext) 
    
    if not content_type:
        raise HTTPException(status_code=422, detail=f"Unsupported file type: {file_ext}")

    # Create uploads directory if it doesn't exist
    upload_dir = "uploads"
    os.makedirs(upload_dir, exist_ok=True)

    # Create a unique filename
    unique_filename = f"chat_{chat_id}_{file.filename}"
    file_path = os.path.join(upload_dir, unique_filename)

    # Save the file
    try:
        contents = await file.read()
        with open(file_path, "wb") as f:
            f.write(contents)
        
        # Convert file to text and print
        extracted_text = convert_file_to_text(file_path, content_type)
        print(f"\n=== Extracted text from {file.filename} ===")
        print(extracted_text)
        print("=" * 50)
        
        # Generate summary of the extracted text with error handling
        try:
            if not extracted_text or len(extracted_text.strip()) == 0:
                summary = "The file appears to be empty or could not be processed."
            else:
                print(f"[DEBUG] Attempting to summarize text of length: {len(extracted_text)}")
                summary = get_text_summary(extracted_text)
                if not summary:
                    summary = "Could not generate summary. The text might be too short or in an unsupported format."
        except Exception as e:
            print(f"[ERROR] Summarization error: {str(e)}")
            summary = "An error occurred while generating the summary."
        
        # Create a message for the file upload
        file_message = Message(
            chat_id=chat_id,
            role="user",
            content=f"📎 Uploaded file: {file.filename}"
        )
        db.add(file_message)

        # Create a message for the summary
        summary_message = Message(
            chat_id=chat_id,
            role="assistant",
            content=f"Here's a summary of {file.filename}:\n\n{summary}"
        )
        db.add(summary_message)
        db.commit()
        
        return {"filename": unique_filename, "summary": summary}
    except Exception as e:
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        await file.close()
